import streamlit as st
from jamaibase import JamAI, protocol as p
import os
from docx import Document
from io import BytesIO
import random
import string
from PyPDF2 import PdfReader
from dotenv import load_dotenv
import os

# Load environment variables from the .env file
load_dotenv()

# Retrieve the values from the environment
api_key = os.getenv("API_KEY")
project_id = os.getenv("PROJECT_ID")

# Initialize the JamAI object
jamai = JamAI(api_key=api_key, project_id=project_id)

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf = PdfReader(pdf_file)
    text = ""
    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


# Function to generate a random filename
def generate_random_filename(extension=".docx"):
    random_str = ''.join(random.choices(string.ascii_letters + string.digits, k=10))
    return f"final_report_{random_str}{extension}"


# Set up the Streamlit app
st.set_page_config(page_title="RelaxingGO", page_icon="🌍")
st.title("🌟 RelaxingGO 🌟 \n- Your Personalized Travel Assistant")

# Custom CSS to style the UI
st.markdown(
    """
    <style>
    body {
        background-color: #e0f7fa;
        color: #004d40;
    }
    .generated-output {
        background-color: #b2dfdb;
        padding: 15px;
        border-radius: 10px;
        margin-top: 20px;
        box-shadow: 0px 4px 10px rgba(0, 0, 0, 0.5);
        color: #004d40;
    }
    .generated-output h4 {
        color: #00796b;
    }
    .stButton>button {
        background-color: #00796b;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        cursor: pointer;
    }
    .stButton>button:hover {
        background-color: #004d40;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Inputs
# First section: Upload Tour Package
st.header("📤 Upload Your Tour Package")
with st.container():
    tour_package = st.file_uploader("📄 Tour Package in PDF format", type="pdf")
    if tour_package:
        # Handle file upload here, for example, extract text or preview the file
        st.success("Tour package uploaded successfully!")

# Second section: Provide Preferences
st.header("🔎 Provide Your Preferences")
with st.container():
    destination = st.text_input("🌍 Desired Destination (e.g., Paris, Tokyo...)")
    travel_dates = st.text_input("📅 Travel Dates (e.g., Dec 20-25, 2024)")
    budget = st.slider("💰 Budget (in MYR)", 500, 10000, 2000)
    activities = st.multiselect(
        "🎯 Select Activities",
        options=[
        "Sightseeing", "Food & Dining", "Adventure Sports", "Relaxation", 
        "Shopping", "Cultural Experiences", "Hiking", "Nightlife", 
        "Photography", "Wellness Retreats", "Historical Tours"
        ]
    )
    allergies = st.text_input("🍴 Any food allergies? (e.g., peanuts, shellfish, beef...)")
    additional_notes = st.text_area("📝 Additional Notes or Preferences")

preferences = (
    "Desired Destination: " + destination +
    ", Travel Dates: " + travel_dates +
    ", Budget in MYR: " + str(budget) +
    ", Activities: " + ", ".join(activities) +
    ", Food Allergies: " + allergies +
    ", Additional notes: " + additional_notes
)

# Action to process inputs
if st.button("🚀 Generate Recommendations", use_container_width=True):
    # Check if all the necessary information is provided
    if tour_package and destination and travel_dates and activities and allergies:
        try:
            # Use st.spinner to show a "Generating..." message
            with st.spinner("Generating recommendations..."):

                # Extract text from the uploaded PDF file
                pdf_text = extract_text_from_pdf(tour_package)

                # Check if the extracted text is null or empty
                if not pdf_text.strip():  # `.strip()` ensures no whitespace-only strings pass
                    st.error("⚠️ The uploaded PDF contains no readable text. Please upload a valid PDF file.")
                else:
                    completion = jamai.add_table_rows(
                        "action",
                        p.RowAddRequest(
                            table_id="RelaxingGO",
                            data=[{"tour_package": pdf_text, "preferences": preferences}],
                            stream=False
                        )
                    )

                    # Display the output generated in the columns
                    if completion.rows:
                        output_row = completion.rows[0].columns
                        summary = output_row.get("summary")
                        attraction = output_row.get("attraction")
                        food = output_row.get("food")
                        season = output_row.get("season")
                        budget = output_row.get("budget")
                        logistics = output_row.get("logistics")
                        recommendation = output_row.get("recommendation")

                        st.subheader("✨ Generated Output")
                        st.markdown(
                            f"""
                            <div class="generated-output">
                                <h4>📝 Summary: </h4> <p>{summary.text if summary else 'N/A'}</p>
                                <h4>🏞️ Attraction: </h4> <p>{attraction.text if attraction else 'N/A'}</p>
                                <h4>🍴 Food:</h4> <p>{food.text if food else 'N/A'}</p>
                                <h4>🌦️ Season: </h4> <p>{season.text if season else 'N/A'}</p>
                                <h4>💰 Budget: </h4> <p>{budget.text if budget else 'N/A'}</p>
                                <h4>🚗 Logistics: </h4> <p>{logistics.text if logistics else 'N/A'}</p>
                                <h4>🧳 Recommendation: </h4> <p>{recommendation.text if recommendation else 'N/A'}</p>
                            </div>
                            """,
                            unsafe_allow_html=True
                        )

                        # Download the final report as a .docx file
                        with st.container():
                            st.subheader("📥 Download Final Report")
                            doc = Document()
                            doc.add_heading("Ratings", level=1)
                            # Sections for each output
                            doc.add_heading("Summary", level=2)
                            doc.add_paragraph(summary.text if summary else 'N/A')
                            doc.add_heading("Attraction", level=2)
                            doc.add_paragraph(attraction.text if attraction else 'N/A')
                            doc.add_heading("Food", level=2)
                            doc.add_paragraph(food.text if food else 'N/A')
                            doc.add_heading("Season", level=2)
                            doc.add_paragraph(season.text if season else 'N/A')
                            doc.add_heading("Budget", level=2)
                            doc.add_paragraph(budget.text if budget else 'N/A')
                            doc.add_heading("Logistics", level=2)
                            doc.add_paragraph(logistics.text if logistics else 'N/A')
                            doc.add_heading("Recommendation", level=2)
                            doc.add_paragraph(recommendation.text if recommendation else 'N/A')

                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            st.download_button(
                                label="📄 Download Final Report as .docx",
                                data=buffer,
                                file_name=generate_random_filename(),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    else:
                        st.error("⚠️ Failed to get a response. Please try again.")
        except Exception as e:
            st.error(f"❌ An error occurred: {e}")
    else:
        st.warning("⚠️ Please fill in all the required information.")